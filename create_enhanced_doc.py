#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Create enhanced technical Word document"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set default styles
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)

def add_title(text, level=1):
    if len(doc.paragraphs) > 1:
        doc.add_page_break()
    
    h = doc.add_heading(text, level=level)
    h.runs[0].font.name = 'Arial'
    h.runs[0].font.size = Pt(18 if level == 1 else 16 if level == 2 else 14)
    h.runs[0].font.bold = True
    h.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_para(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.15
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
    
    return table

# Title page
title = doc.add_heading('ניתוח טכני מעמיק', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(28)

doc.add_paragraph()

subtitle = doc.add_paragraph('I figured out the best way to run OpenClaw')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(18)
subtitle.runs[0].font.bold = True

meta = doc.add_paragraph('Matthew Berman | YouTube Technical Analysis')
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# Executive Summary
add_title('תקציר מנהלים', 1)

add_para('סרטון זה מציג מדריך מקיף להתקנה ושימוש ב-OpenClaw.')
add_para('Matthew Berman מדגים כיצד OpenClaw הפך לכלי חיוני בעבודתו היומיומית.')

# Architecture
add_title('ארכיטקטורת המערכת', 1)
add_para('OpenClaw בנוי כמסגרת מודולרית המורכבת משכבות מרובות.')

components = [
    ('Gateway', 'WebSocket Server', 'נקודת כניסה מרכזית'),
    ('Agent Engine', 'AI Processing', 'מנוע הבינה המלאכותית'),
    ('Memory System', 'File + Vector DB', 'ניהול זיכרון'),
    ('Skills Manager', 'Plugin System', 'ניהול Skills'),
]

add_table(['רכיב', 'טכנולוגיה', 'תפקיד'], components)

# Tools
doc.add_page_break()
add_title('כלים וטכנולוגיות', 1)

tools = [
    ('OpenClaw', 'Core Platform', 'Open Source', 'github.com/openclaw'),
    ('Hostinger VPS', 'Cloud Hosting', 'Commercial', 'hostinger.com'),
    ('Telegram', 'Messaging', 'Free', 'telegram.org'),
    ('Claude', 'LLM', 'Paid', 'anthropic.com'),
    ('Gmail API', 'Email', 'Free', 'Google'),
]

add_table(['כלי', 'קטגוריה', 'רישוי', 'מקור'], tools)

# Installation
doc.add_page_break()
add_title('התקנה ותצורה', 1)

add_para('שלב 1: רכישת VPS')
add_para('1. גש ל-hostinger.com/matthewb')
add_para('2. בחר תוכנית VPS')
add_para('3. הזן קוד: MatthewB')
add_para('4. השלם תשלום')

add_para('שלב 2: תצורה')
add_para('1. הזן API keys')
add_para('2. לחץ Deploy')
add_para('3. חבר Telegram')

# Commands
doc.add_page_break()
add_title('מדריך פקודות', 1)

commands = [
    ('openclaw status', 'הצג סטטוס', 'משתמש רגיל'),
    ('openclaw init', 'התחל onboarding', 'התקנה'),
    ('openclaw restart', 'אתחל Gateway', 'אדמין'),
    ('openclaw mcp list', 'הצג MCP servers', 'משתמש'),
]

add_table(['פקודה', 'תיאור', 'הרשאות'], commands)

# Configuration files
doc.add_page_break()
add_title('קבצי תצורה', 1)

files = [
    ('SOUL.md', 'Personality', 'קריטי', 'עדכון ידני'),
    ('IDENTITY.md', 'Basic Info', 'חובה', 'התקנה'),
    ('MEMORY.md', 'Long-term Memory', 'קריטי', 'תקופתי'),
    ('USER.md', 'User Info', 'חשוב', 'שוטף'),
]

add_table(['קובץ', 'מטרה', 'חשיבות', 'תדירות'], files)

# Skills
doc.add_page_break()
add_title('מערכת Skills', 1)

add_para('Skills הם מודולים עצמאיים המרחיבים את יכולות המערכת.')

skills = [
    ('weather', 'מזג אוויר', 'wttr.in', 'מובנה'),
    ('browser', 'שליטה בדפדפן', 'Playwright', 'מובנה'),
    ('google-workspace', 'Gmail/Calendar', 'OAuth', 'מובנה'),
]

add_table(['Skill', 'תיאור', 'טכנולוגיה', 'סטטוס'], skills)

# Save
output_path = 'youtube_chen_output/3GrG-dOmrLU_TECHNICAL_FULL.docx'
doc.save(output_path)

print(f"✅ Document created: {output_path}")
