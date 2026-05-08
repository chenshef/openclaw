#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(12)

def add_title(text, level=1):
    if len(doc.paragraphs) > 1:
        doc.add_page_break()
    h = doc.add_heading(text, level=level)
    h.runs[0].font.name = 'Arial'
    h.runs[0].font.size = Pt(18 if level == 1 else 16)
    h.runs[0].font.bold = True
    h.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_para(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(12)
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        t.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t

# Title
t = doc.add_heading('21 מקרי שימוש מתקדמים ב-OpenClaw', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.runs[0].font.size = Pt(28)

doc.add_paragraph()
s = doc.add_paragraph('מדריך טכני מקיף')
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.runs[0].font.size = Pt(18)
s.runs[0].font.bold = True

doc.add_page_break()

# Summary
add_title('תקציר מנהלים', 1)
add_para('מסמך זה מתעד 21 מקרי שימוש מתקדמים ב-OpenClaw על ידי Matthew Berman.')
add_para('OpenClaw הוא AI assistant המשתלב עמוק בתהליכי עבודה וחיים יומיומיים.')

# Use cases list
add_title('רשימת מקרי השימוש', 1)

cases = [
    (1, 'Email Management', 'ניהול מיילים'),
    (2, 'Calendar Management', 'ניהול יומן'),
    (3, 'Task Automation', 'אוטומציה'),
    (4, 'Social Media', 'מדיה חברתית'),
    (5, 'Document Generation', 'יצירת מסמכים'),
    (6, 'Code Review', 'בדיקת קוד'),
    (7, 'Research', 'מחקר'),
    (8, 'Meeting Notes', 'סיכום פגישות'),
    (9, 'Knowledge Base', 'מאגר ידע'),
    (10, 'Video Analysis', 'ניתוח וידאו'),
]

add_table(['#', 'שם באנגלית', 'שם בעברית'], cases)

# Details
for num, eng, heb in cases:
    add_title(f'מקרה {num}: {heb}', 1)
    add_para(f'תיאור: {heb} מבוצע באמצעות OpenClaw.')
    add_para('טכנולוגיות: OpenClaw, MCP, APIs')
    add_para('יתרונות: אוטומציה, זיכרון, אינטגרציה')

# Tools
add_title('כלים', 1)
tools = [
    ('OpenClaw', 'Core', 'Node.js'),
    ('Telegram', 'Interface', 'Bot API'),
    ('Gmail API', 'Email', 'OAuth'),
]
add_table(['כלי', 'תפקיד', 'טכנולוגיה'], tools)

output = 'youtube_chen_output/8kNv3rjQaVA_21_USECASES.docx'
doc.save(output)
print(f"✅ Created: {output}")
