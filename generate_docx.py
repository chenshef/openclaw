#!/usr/bin/env python3
"""
Generate Word (DOCX) document from Markdown for OpenClaw Technical Guide
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Style as hyperlink (blue, underlined)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def parse_markdown_to_docx(md_file, docx_file):
    """Convert Markdown to DOCX with formatting."""
    doc = Document()
    
    # Set default styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_code_block = False
    code_lang = ''
    in_table = False
    table_headers = []
    table = None
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip mermaid diagrams
        if line.startswith('```mermaid'):
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            i += 1
            continue
        
        # Code blocks
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lang = line[3:].strip()
                i += 1
                continue
            else:
                in_code_block = False
                i += 1
                continue
        
        if in_code_block:
            p = doc.add_paragraph(style='Normal')
            p_format = p.paragraph_format
            p_format.left_indent = Inches(0.5)
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.runs[0].font.size = Pt(20)
            p.runs[0].font.name = 'Arial'
            p.runs[0].bold = True
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            p.runs[0].font.size = Pt(16)
            p.runs[0].font.name = 'Arial'
            p.runs[0].bold = True
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.name = 'Arial'
            p.runs[0].bold = True
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
            p.runs[0].font.size = Pt(12)
            p.runs[0].font.name = 'Arial'
            p.runs[0].bold = True
        
        # Horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph('_' * 50)
        
        # Tables
        elif '|' in line and line.strip().startswith('|'):
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            
            if not in_table:
                # Start new table
                in_table = True
                table_headers = cells
                # Peek ahead to check for separator line
                if i + 1 < len(lines) and '---' in lines[i + 1]:
                    i += 1  # Skip separator line
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Light Grid Accent 1'
                # Add headers
                hdr_cells = table.rows[0].cells
                for idx, header in enumerate(cells):
                    hdr_cells[idx].text = header
                    # Bold headers
                    for paragraph in hdr_cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(12)
            else:
                # Add row
                if table and len(cells) == len(table_headers):
                    row_cells = table.add_row().cells
                    for idx, cell_text in enumerate(cells):
                        row_cells[idx].text = cell_text
                        for paragraph in row_cells[idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(11)
        else:
            # Check if we just left a table
            if in_table and '|' not in line:
                in_table = False
                table = None
                table_headers = []
            
            # Blockquotes
            if line.startswith('> '):
                p = doc.add_paragraph(line[2:])
                p.paragraph_format.left_indent = Inches(0.5)
                p.runs[0].font.italic = True
                p.runs[0].font.color.rgb = RGBColor(80, 80, 80)
            
            # Lists
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                text = line.strip()[2:]
                p = doc.add_paragraph(text, style='List Bullet')
                p.runs[0].font.size = Pt(12)
            
            elif re.match(r'^\d+\.\s', line.strip()):
                text = re.sub(r'^\d+\.\s', '', line.strip())
                p = doc.add_paragraph(text, style='List Number')
                p.runs[0].font.size = Pt(12)
            
            # Regular paragraphs
            elif line.strip():
                # Handle inline formatting
                p = doc.add_paragraph()
                
                # Process bold, italic, code, links
                text = line
                
                # Links: [text](url)
                link_pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
                parts = []
                last_end = 0
                
                for match in re.finditer(link_pattern, text):
                    # Add text before link
                    if match.start() > last_end:
                        parts.append(('text', text[last_end:match.start()]))
                    # Add link
                    parts.append(('link', match.group(1), match.group(2)))
                    last_end = match.end()
                
                # Add remaining text
                if last_end < len(text):
                    parts.append(('text', text[last_end:]))
                
                # If no links, just add the text with basic formatting
                if not parts:
                    # Handle bold/italic/code
                    segments = []
                    current = ''
                    in_bold = False
                    in_italic = False
                    in_code = False
                    
                    j = 0
                    while j < len(text):
                        if text[j:j+2] == '**':
                            if current:
                                segments.append((current, in_bold, in_italic, in_code))
                                current = ''
                            in_bold = not in_bold
                            j += 2
                        elif text[j] == '*':
                            if current:
                                segments.append((current, in_bold, in_italic, in_code))
                                current = ''
                            in_italic = not in_italic
                            j += 1
                        elif text[j] == '`':
                            if current:
                                segments.append((current, in_bold, in_italic, in_code))
                                current = ''
                            in_code = not in_code
                            j += 1
                        else:
                            current += text[j]
                            j += 1
                    
                    if current:
                        segments.append((current, in_bold, in_italic, in_code))
                    
                    for seg_text, is_bold, is_italic, is_code in segments:
                        run = p.add_run(seg_text)
                        run.font.size = Pt(12)
                        if is_bold:
                            run.font.bold = True
                        if is_italic:
                            run.font.italic = True
                        if is_code:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(200, 0, 0)
                else:
                    # Process parts (text and links)
                    for part in parts:
                        if part[0] == 'text':
                            run = p.add_run(part[1])
                            run.font.size = Pt(12)
                        elif part[0] == 'link':
                            add_hyperlink(p, part[2], part[1])
            else:
                # Empty line
                doc.add_paragraph()
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✓ DOCX document created: {docx_file}")

if __name__ == '__main__':
    parse_markdown_to_docx(
        'OPENCLAW_COMPLETE_TECHNICAL_GUIDE.md',
        'OPENCLAW_COMPLETE_TECHNICAL_GUIDE.docx'
    )
